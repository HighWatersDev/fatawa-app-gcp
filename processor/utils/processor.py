from docx import Document
from pathlib import Path


def collect_audio_data(base_path):
    audio_data = []

    base_path = Path(base_path)

    for author_path in base_path.iterdir():
        if author_path.is_dir():
            for folder_path in author_path.iterdir():
                if folder_path.is_dir():
                    for audio_file in folder_path.glob("*.mp3"):
                        transcription_file = audio_file.stem + "_transcription.txt"
                        transcription_path = folder_path / transcription_file
                        additional_metadata = "..."  # Add logic to get additional metadata

                        if transcription_path.exists():
                            with transcription_path.open('r', encoding='utf-8') as transcription_file:
                                transcription_text = transcription_file.read().strip()

                            audio_data.append((
                                author_path.name,
                                folder_path.name,
                                audio_file.name,
                                transcription_text,
                                additional_metadata
                            ))

    return audio_data


# Example usage
base_directory = "path/to/audio-files"  # Replace with your actual base directory
data = collect_audio_data(base_directory)
print(data)

# Printing the results
for row in data:
    print({
        "author": row[0],
        "folder": row[1],
        "filename": row[2],
        "transcription": row[3],
        "additional_metadata": row[4]
    })


def process_word_file(file_path):
    doc = Document(file_path)

    questions = []
    answers = []

    for i in range(0, len(doc.paragraphs), 2):
        question = doc.paragraphs[i].text.strip()
        answer = doc.paragraphs[i + 1].text.strip()

        questions.append(question)
        answers.append(answer)

    return questions, answers


file_path = "/Users/luqmaan/Projects/fatawa/artifacts/transcriptions/ruhayli/upload1527872508122/1.docx"
questions, answers = process_word_file(file_path)

doc = Document(file_path)

# Printing the results
for q, a in zip(questions, answers):
    print(f"Q: {q}")
    print(f"A: {a}")
    print("-----")